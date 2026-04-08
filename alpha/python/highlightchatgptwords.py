import pandas as pd
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from collections import Counter

# Load the CSV file
url = "https://raw.githubusercontent.com/berenslab/chatgpt-excess-words/main/results/excess_words.csv"
df = pd.read_csv(url)
excess_words = [word.lower() for word in df.iloc[:, 0].tolist()]  # Convert all words to lower case

# Function to create a highlighted run
def create_highlighted_run(paragraph, text):
    run = paragraph.add_run(text)
    highlight = OxmlElement("w:highlight")
    highlight.set(qn("w:val"), "blue")  # Highlight background color
    run._element.get_or_add_rPr().append(highlight)
    return run

# Open the Word document
doc_path = "C:\\temp\\2024 Unafraid.docx"  # Replace with your .docx file path
doc = Document(doc_path)

# Dictionary to count word occurrences
word_count = Counter()

# Highlight excess words and count occurrences
for paragraph in doc.paragraphs:
    new_paragraph_text = []
    for run in paragraph.runs:
        # Find word-like objects, including punctuation
        word_like_objects = re.findall(r"\S+", run.text)
        for obj in word_like_objects:
            words_in_obj = re.findall(r'\w+|[^\w\s]', obj, re.UNICODE)  # Further split to handle punctuation
            if any(word.lower() in excess_words for word in words_in_obj):
                word_count.update([word.lower() for word in words_in_obj if word.lower() in excess_words])
                new_paragraph_text.append((obj, True))  # Mark object for highlighting
                print(f"Encountered word-like object: {obj}")
            else:
                new_paragraph_text.append((obj, False))  # Mark object without highlighting

    # Reconstruct paragraph with highlighted word-like objects
    paragraph.clear()
    for obj, highlight in new_paragraph_text:
        if highlight:
            create_highlighted_run(paragraph, obj)
        else:
            paragraph.add_run(obj)
        paragraph.add_run(' ')  # Add a space after each word-like object

# Save the modified document
doc.save("C:\\Temp\\highlighted_document.docx")

# Output the word count to a CSV file
word_count_df = pd.DataFrame(word_count.items(), columns=['Word', 'Count'])
word_count_df.to_csv("C:\\Temp\\word_count.csv", index=False)

print("Highlighting complete. The document has been saved as 'highlighted_document.docx'.")
print("Word count has been saved as 'word_count.csv'.")
